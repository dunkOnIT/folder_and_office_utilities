from pathlib import Path
from typing import Union

import aspose.words as aw
from docx import Document
from python_docx_replace.docx_replace import docx_replace

doc = aw.Document()

class WordDoc:
    def __init__(self, path):
        self.path = path
        self.document = Document(path)

    def find_and_replace_in_document(self, replacement_mapping: dict[str, str]):
        """For each key/value pair in the given replacement mapping,
        iterate through the document and replace each key with its value."""

        for replacement_target in replacement_mapping.keys():
            print(replacement_target)
            print(replacement_mapping[replacement_target])

            for p in self.document.paragraphs:
                for run in p.runs:
                    if replacement_target in run.text:
                        run.text = run.text.replace(replacement_target, replacement_mapping[replacement_target])

    def save(self, filename: str, path: Union[str, Path, None] = None) -> None:
        if path is None:
            save_path = Path(filename)
        else:
            if not Path(path).exists():
                Path(path).mkdir()  # Create the target save directory if it doesn't already exist

            save_path = Path(path).joinpath(filename)

        self.document.save(save_path)
